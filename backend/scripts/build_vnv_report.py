"""
Parses the master Markdown V&V report into a simple block structure, then
renders that SAME block list into both DOCX (python-docx) and PDF
(reportlab), so all three output formats stay content-identical - no
separate hand-authored content per format.
"""
from __future__ import annotations
import re
from pathlib import Path

REPORT_DIR = Path(__file__).resolve().parents[1] / "src" / "evaluation" / "vnv_report"
MD_PATH = REPORT_DIR / "LOCUS_AI_RAG_VnV_Report.md"


def parse_markdown(text: str) -> list[dict]:
    lines = text.split("\n")
    blocks = []
    i = 0
    in_frontmatter = False
    while i < len(lines):
        line = lines[i]

        if line.strip() == "---" and i == 0:
            in_frontmatter = True
            i += 1
            continue
        if in_frontmatter:
            if line.strip() == "---":
                in_frontmatter = False
            i += 1
            continue

        if line.strip() == "---":
            blocks.append({"type": "hr"})
            i += 1
            continue

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append({"type": "heading", "level": level, "text": heading_match.group(2).strip()})
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not re.match(r"^\|[\s:\-|]+\|$", lines[i].strip()):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    table_rows.append(cells)
                i += 1
            blocks.append({"type": "table", "rows": table_rows})
            continue

        if re.match(r"^\s*[-*]\s+", line):
            list_items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                list_items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            blocks.append({"type": "list", "items": list_items})
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            list_items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                list_items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append({"type": "numlist", "items": list_items})
            continue

        if line.strip() == "":
            i += 1
            continue

        # plain paragraph (collect consecutive non-blank, non-special lines)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() != "" and not re.match(
            r"^(#{1,3}\s|```|\|.*\||\s*[-*]\s|\s*\d+\.\s|---\s*$)", lines[i]
        ):
            para_lines.append(lines[i])
            i += 1
        blocks.append({"type": "para", "text": " ".join(para_lines)})

    return blocks


def _split_bold(text: str) -> list[tuple[str, bool, bool]]:
    """Split text on **bold** and `code` markers, return [(text, is_bold, is_code), ...]."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    out = []
    for p in parts:
        if p.startswith("**") and p.endswith("**"):
            out.append((p[2:-2], True, False))
        elif p.startswith("`") and p.endswith("`") and len(p) > 1:
            out.append((p[1:-1], False, True))
        elif p:
            out.append((p, False, False))
    return out


def build_docx(blocks: list[dict], out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    def shade_cell(cell, color_hex):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color_hex)
        tc_pr.append(shd)

    first_heading_done = False
    for b in blocks:
        if b["type"] == "heading":
            if b["level"] == 1 and not first_heading_done:
                title = doc.add_heading(b["text"], level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_heading_done = True
            else:
                doc.add_heading(b["text"], level=min(b["level"], 3))
        elif b["type"] == "para":
            p = doc.add_paragraph()
            for text, bold, code in _split_bold(b["text"]):
                run = p.add_run(text)
                run.bold = bold
                if code:
                    run.font.name = "Courier New"
        elif b["type"] == "list":
            for item in b["items"]:
                p = doc.add_paragraph(style="List Bullet")
                for text, bold, code in _split_bold(item):
                    run = p.add_run(text)
                    run.bold = bold
                    if code:
                        run.font.name = "Courier New"
        elif b["type"] == "numlist":
            for item in b["items"]:
                p = doc.add_paragraph(style="List Number")
                for text, bold, code in _split_bold(item):
                    run = p.add_run(text)
                    run.bold = bold
                    if code:
                        run.font.name = "Courier New"
        elif b["type"] == "code":
            p = doc.add_paragraph()
            run = p.add_run(b["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            p.paragraph_format.left_indent = Inches(0.25)
        elif b["type"] == "table":
            rows = b["rows"]
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx >= len(table.rows[r_idx].cells):
                        continue
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    for text, bold, code in _split_bold(cell_text):
                        run = p.add_run(text)
                        run.bold = bold or r_idx == 0
                        run.font.size = Pt(9)
                        if code:
                            run.font.name = "Courier New"
                    if r_idx == 0:
                        shade_cell(cell, "D9E2F3")
        elif b["type"] == "hr":
            doc.add_page_break()

    doc.save(str(out_path))


def build_pdf(blocks: list[dict], out_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        ListFlowable, ListItem, PageBreak, Preformatted,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCustom", parent=styles["Title"], fontSize=20, spaceAfter=18)
    h1 = ParagraphStyle("H1Custom", parent=styles["Heading1"], fontSize=15, spaceBefore=16, spaceAfter=8)
    h2 = ParagraphStyle("H2Custom", parent=styles["Heading2"], fontSize=12.5, spaceBefore=12, spaceAfter=6)
    h3 = ParagraphStyle("H3Custom", parent=styles["Heading3"], fontSize=11, spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("BodyCustom", parent=styles["BodyText"], fontSize=9.5, leading=13, spaceAfter=8)
    code_style = ParagraphStyle("CodeCustom", parent=styles["Code"], fontSize=7.5, leading=9.5,
                                 backColor=colors.whitesmoke, borderPadding=6)

    def inline_to_html(text):
        text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"`([^`]+)`", r'<font face="Courier">\1</font>', text)
        return text

    doc = SimpleDocTemplate(str(out_path), pagesize=LETTER,
                             topMargin=0.75 * inch, bottomMargin=0.75 * inch,
                             leftMargin=0.75 * inch, rightMargin=0.75 * inch)
    story = []
    first_heading_done = False

    for b in blocks:
        if b["type"] == "heading":
            if b["level"] == 1 and not first_heading_done:
                story.append(Paragraph(b["text"], title_style))
                first_heading_done = True
            elif b["level"] == 1:
                story.append(Paragraph(b["text"], h1))
            elif b["level"] == 2:
                story.append(Paragraph(b["text"], h2))
            else:
                story.append(Paragraph(b["text"], h3))
        elif b["type"] == "para":
            story.append(Paragraph(inline_to_html(b["text"]), body))
        elif b["type"] in ("list", "numlist"):
            items = [ListItem(Paragraph(inline_to_html(t), body)) for t in b["items"]]
            story.append(ListFlowable(items, bulletType="bullet" if b["type"] == "list" else "1"))
            story.append(Spacer(1, 6))
        elif b["type"] == "code":
            story.append(Preformatted(b["text"], code_style))
            story.append(Spacer(1, 8))
        elif b["type"] == "table":
            rows = b["rows"]
            if not rows:
                continue
            data = [[Paragraph(inline_to_html(c), ParagraphStyle("cell", parent=body, fontSize=8, leading=10))
                     for c in row] for row in rows]
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))
        elif b["type"] == "hr":
            story.append(PageBreak())

    doc.build(story)


def main():
    text = MD_PATH.read_text()
    blocks = parse_markdown(text)
    print(f"Parsed {len(blocks)} blocks from {MD_PATH.name}")

    docx_path = REPORT_DIR / "LOCUS_AI_RAG_VnV_Report.docx"
    pdf_path = REPORT_DIR / "LOCUS_AI_RAG_VnV_Report.pdf"

    build_docx(blocks, docx_path)
    print(f"Wrote {docx_path}")

    build_pdf(blocks, pdf_path)
    print(f"Wrote {pdf_path}")


if __name__ == "__main__":
    main()
